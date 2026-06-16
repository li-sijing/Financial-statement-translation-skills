#!/tmp/docx_venv/bin/python3
"""Reusable helpers for batch translation scripts.
Copy this header into each batch script, then call replace_para() and fmt_table().

Usage:
    from docx import Document
    doc = Document("path.docx")
    replace_para(doc.paragraphs[5], "English text", bold=True)
    fmt_table(doc.tables[0])
    doc.save("path.docx")
"""
import sys, re
sys.path.insert(0, '/tmp/docx_venv/lib/python3.14/site-packages')
from docx.shared import Pt
from docx.oxml.ns import qn

TN = 'Times New Roman'
AN = 'Arial Narrow'
FS = Pt(10.5)

def set_run(run, text, bold=False, font_name=TN):
    run.text = text
    run.font.name = font_name
    run.font.size = FS
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def replace_para(para, text, bold=False):
    """Replace paragraph text with English, applying Times New Roman 10.5pt."""
    for r in para.runs:
        r.text = ''
    run = para.add_run(text)
    set_run(run, text, bold=bold)

def replace_cell(cell, text, font_name=TN, bold=False):
    """Replace table cell text with specified font."""
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(text)
    set_run(run, text, bold=bold, font_name=font_name)

def is_numeric(s):
    """Check if cell content is a number (needs Arial Narrow)."""
    s = s.replace(',', '').replace('.', '').replace('-', '').replace('%', '').replace(' ', '').strip()
    return bool(s) and all(c.isdigit() or c in '.,-%' for c in s)

def fmt_table(t):
    """Auto-apply Arial Narrow to numeric cells in a table."""
    for row in t.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            if not txt:
                continue
            if is_numeric(txt):
                replace_cell(cell, txt, AN)
